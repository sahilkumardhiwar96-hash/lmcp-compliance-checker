# -*- coding: utf-8 -*-
"""Editable report exports: CSV (raw data) and Word/DOCX (formatted, bilingual)."""

import io
import csv
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from translations import PDF_STRINGS, get_field_label


def _field_key(item, field_labels_en):
    for key, label in field_labels_en.items():
        if label == item["label"]:
            return key
    return None


def _set_complex_script_font(run, font_name):
    """Ensure the complex-script (e.g. Devanagari) font is set explicitly,
    so Word shapes Hindi text correctly regardless of the default document font."""
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:cs"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def generate_csv_report(found, missing, score, filename, location_name=None):
    """Simple raw-data CSV export — English column headers, useful for spreadsheets/analysis."""
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow(["Product/file", filename])
    writer.writerow(["Scan date/time", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
    writer.writerow(["Compliance score (%)", score])
    if location_name:
        writer.writerow(["Location", location_name])
    writer.writerow([])
    writer.writerow(["Status", "Declaration", "Value", "Legal reference", "Notes"])
    for f in found:
        writer.writerow(["Found", f["label"], f.get("value") or "", f["legal_ref"], ""])
    for m in missing:
        writer.writerow(["Missing", m["label"], "", m["legal_ref"], m.get("note", "")])
    return io.BytesIO(buffer.getvalue().encode("utf-8-sig"))  # BOM so Excel opens Unicode correctly


def generate_docx_report(image_bytes, found, missing, score, filename, field_labels_en, lang="en", location_name=None):
    """Formatted, editable Word report — English or Hindi. Word shapes Devanagari
    correctly on its own, so no special font-embedding is needed (unlike PDF)."""
    s = PDF_STRINGS[lang]
    font_name = "Nirmala UI" if lang == "hi" else "Calibri"  # Nirmala UI ships with Windows and renders Devanagari well

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:cs"), font_name)

    title = doc.add_heading(s["title"], level=1)
    for run in title.runs:
        _set_complex_script_font(run, font_name)

    sub = doc.add_paragraph(s["subtitle"])
    sub.runs[0].font.size = Pt(9.5)
    sub.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    _set_complex_script_font(sub.runs[0], font_name)

    meta_lines = [
        f"{s['product_file']}: {filename}",
        f"{s['scan_time']}: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"{s['score']}: {score}%",
    ]
    if location_name:
        meta_lines.append(f"{s['location']}: {location_name}")
    for line in meta_lines:
        p = doc.add_paragraph(line)
        _set_complex_script_font(p.runs[0], font_name)

    if image_bytes:
        h = doc.add_heading(s["photo_evidence"], level=2)
        for run in h.runs:
            _set_complex_script_font(run, font_name)
        try:
            doc.add_picture(io.BytesIO(image_bytes), width=Inches(3))
        except Exception:
            pass

    # ---- Found declarations table ----
    h = doc.add_heading(s["found_heading"], level=2)
    for run in h.runs:
        _set_complex_script_font(run, font_name)

    if found:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        headers = [s["col_declaration"], s["col_value"], s["col_legal_ref"]]
        for i, htext in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(htext)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_complex_script_font(run, font_name)
            _shade_cell(cell, "2E7D32")
        for f in found:
            row = table.add_row().cells
            values = [
                get_field_label(_field_key(f, field_labels_en), field_labels_en, lang),
                str(f.get("value") or "-"),
                f["legal_ref"],
            ]
            for i, val in enumerate(values):
                row[i].text = ""
                run = row[i].paragraphs[0].add_run(val)
                _set_complex_script_font(run, font_name)
    else:
        p = doc.add_paragraph(s["none_found"])
        _set_complex_script_font(p.runs[0], font_name)

    # ---- Missing declarations table ----
    h = doc.add_heading(s["missing_heading"], level=2)
    for run in h.runs:
        _set_complex_script_font(run, font_name)

    if missing:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        headers = [s["col_declaration"], s["col_legal_ref_notes"]]
        for i, htext in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(htext)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_complex_script_font(run, font_name)
            _shade_cell(cell, "C62828")
        for m in missing:
            ref_text = m["legal_ref"]
            if m.get("note"):
                ref_text += f" — {s['note_prefix']}: {m['note']}"
            row = table.add_row().cells
            values = [get_field_label(_field_key(m, field_labels_en), field_labels_en, lang), ref_text]
            for i, val in enumerate(values):
                row[i].text = ""
                run = row[i].paragraphs[0].add_run(val)
                _set_complex_script_font(run, font_name)

        p = doc.add_paragraph(s["penalty_note"])
        p.runs[0].font.size = Pt(8.5)
        p.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        _set_complex_script_font(p.runs[0], font_name)
    else:
        p = doc.add_paragraph(s["no_violations"])
        _set_complex_script_font(p.runs[0], font_name)

    for text in [s["font_disclaimer"], s["footer_disclaimer"]]:
        p = doc.add_paragraph(text)
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        _set_complex_script_font(p.runs[0], font_name)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
